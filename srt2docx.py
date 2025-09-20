# In this script we are accepting the srt file from user having speaker name and its dialogue and if dialogues of the same speaker is there continously then we are merging those continious dialogues into one entry.

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

def validator(inp_dir):
    return os.path.isdir(inp_dir)

def extract_speaker_dialogues_srt(srt_path):
    with open(srt_path, 'r', encoding='utf-8') as f:
        lines = [line.strip() for line in f if line.strip()]

    extracted = []
    current_speaker = None
    dialogue_buffer = []

    for line in lines:
        if line.isdigit():  # skipping the segment number
            continue
        if "-->" in line:  # skipping the timestamp
            continue

        # checking for speaker name present in "[]"
        if line.startswith('[') and ']' in line:
            end_idx = line.find(']')
            speaker_tag = line[1:end_idx].strip() # Fetching the speaker name
            dialogue_text = line[end_idx+1:].strip() # Its dialogue

            # If speaker changes, saving previous speaker's dialogue
            if speaker_tag != current_speaker:
                if current_speaker and dialogue_buffer:
                    extracted.append((current_speaker, ' '.join(dialogue_buffer)))
                current_speaker = speaker_tag # Chnagin the current speaker
                dialogue_buffer = [dialogue_text] if dialogue_text else []
            else:
            
                dialogue_buffer.append(dialogue_text)
        else:
            #if speaker name is not detected means the speaker is same.
            if current_speaker:
                dialogue_buffer.append(line)
    if current_speaker and dialogue_buffer:
        extracted.append((current_speaker, ' '.join(dialogue_buffer)))

    return extracted # returning the extracted dialogues
# Function for formatting the extracted content
def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)

def is_tamil(text):
    return any('\u0B80' <= char <= '\u0BFF' for char in text)

def write_structured_doc(extracted_data, output_path):
    doc = Document()

    if not extracted_data:
        doc.add_paragraph(" No speaker-tagged dialogue found.")
        doc.save(output_path)
        return

    for speaker, dialogue in extracted_data:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Pt(80)
        table.columns[1].width = Pt(400)
        remove_table_borders(table)

        # Speaker cell
        speaker_cell = table.cell(0, 0)
        speaker_para = speaker_cell.paragraphs[0]
        speaker_run = speaker_para.add_run(speaker)
        speaker_run.bold = True

        # Dialogue cell
        dialogue_cell = table.cell(0, 1)
        dialogue_para = dialogue_cell.paragraphs[0]
        run = dialogue_para.add_run(dialogue)
        if is_tamil(dialogue):
            run.font.name = 'Arial Unicode MS'

    doc.save(output_path)

if __name__=="__main__":
    try:
        input_dir = input("Please enter directory containing .srt files: ").strip()
        if not validator(input_dir):
            raise Exception()
    except Exception:
        print("Please provide a valid directory!")
    else:
        output_folder = os.path.join(input_dir, "Srt to docx output")
        os.makedirs(output_folder, exist_ok=True)

        for each_file in os.listdir(input_dir):
            if each_file.lower().endswith(".srt"):
                srt_path = os.path.join(input_dir, each_file)
                docx_filename=each_file.replace(".srt",".docx")
                output_docx_path = os.path.join(output_folder, docx_filename)

                extracted = extract_speaker_dialogues_srt(srt_path)
                if not extracted:
                    print(f"No speaker-tagged lines found in: {each_file}")
                else:
                    print(f"Parsed {len(extracted)} speaker segments from {each_file}")

                write_structured_doc(extracted, output_docx_path)
                print(f"Saved: {output_docx_path}")
            else:
                print(f"Skipping non-srt file: {each_file}")
